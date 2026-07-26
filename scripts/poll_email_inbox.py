#!/usr/bin/env python3
"""
LedgerX inbound-email poller
────────────────────────────
Runs on the VPS as a cron job (every 5 minutes).
Connects to the receipts@90ten.life IMAP mailbox, fetches unseen messages,
and forwards each one to the Supabase inbound-email edge function.

Cron entry (crontab -e):
    */5 * * * * set -a && . /opt/ledgerx/env && set +a && /opt/ledgerx/venv/bin/python3 /opt/ledgerx/poll_email_inbox.py >> /var/log/ledgerx_email.log 2>&1
    (runs inside a venv, NOT system python3 — that's where weasyprint and,
    as of v13.9, PyMuPDF live; installing to system pip won't reach cron)

Configuration: edit the CONFIG block below, or set the equivalent
environment variables (env vars take precedence).
"""

import imaplib
import email
import email.policy
import base64
import html
import io
import json
import os
import sys
import urllib.request
import urllib.error
from datetime import datetime

# weasyprint is used to render HTML email bodies to PDF when no
# real attachment is present. The Python package import is light;
# rendering requires system libs (Pango, Cairo, etc.) which may
# fail at write_pdf() time even if the import succeeds.
WEASYPRINT_AVAILABLE = False
WEASYPRINT_VERSION = None
WEASYPRINT_IMPORT_ERROR = None
try:
    import weasyprint
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
    WEASYPRINT_VERSION = weasyprint.__version__
except Exception as _ex:
    WEASYPRINT_IMPORT_ERROR = str(_ex)

# PyMuPDF rasterizes PDFs to PNG so they can actually be OCR'd.
#
# Why this exists: OpenAI's vision endpoint accepts JPEG/PNG/WEBP/GIF only —
# it rejects a `data:application/pdf;base64,...` URL outright. The edge
# function was handing it PDFs anyway and swallowing the 400, so every
# PDF-borne receipt landed with `prefilled = {}` and the user got a blank
# form. Retailer receipts are overwhelmingly PDFs (attached, or rendered
# from an HTML body below), so in practice that was most forwarded mail.
#
# Rasterizing here rather than in the edge function keeps the Deno side
# simple and means the fix works against the CURRENTLY DEPLOYED function
# without touching it — we insert the PNG ahead of the PDF, and its
# "first OCR-compatible attachment" scan picks the PNG up on its own.
#
# Guarded like weasyprint: a missing wheel degrades to today's behaviour
# rather than taking the whole poller down.
PYMUPDF_AVAILABLE = False
PYMUPDF_IMPORT_ERROR = None
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except Exception as _ex:
    PYMUPDF_IMPORT_ERROR = str(_ex)

# ── Configuration ─────────────────────────────────────────────────────────────
# Fill these in once you have the Hostinger mailbox credentials.
# You can also set them as environment variables on the VPS.

CONFIG = {
    # Hostinger IMAP settings for receipts@90ten.life
    "IMAP_HOST":     os.environ.get("LEDGERX_IMAP_HOST",     "imap.hostinger.com"),
    "IMAP_PORT":     int(os.environ.get("LEDGERX_IMAP_PORT", "993")),
    "IMAP_USER":     os.environ.get("LEDGERX_IMAP_USER",     "receipts@90ten.life"),
    "IMAP_PASSWORD": os.environ.get("LEDGERX_IMAP_PASSWORD", "REPLACE_WITH_MAILBOX_PASSWORD"),

    # Supabase edge function endpoint
    "FUNCTION_URL":  os.environ.get("LEDGERX_FUNCTION_URL",
                                    "https://bkxccrbfjoqtxbtekrgw.supabase.co/functions/v1/inbound-email"),

    # Shared secret — must match INBOUND_EMAIL_SECRET in Supabase secrets
    "INBOUND_SECRET": os.environ.get("LEDGERX_INBOUND_SECRET", "REPLACE_WITH_SHARED_SECRET"),

    # Max attachment size to upload (bytes). Larger attachments are skipped.
    "MAX_ATTACH_BYTES": 10 * 1024 * 1024,  # 10 MB
}

# Allowed attachment MIME types
ALLOWED_TYPES = {
    "image/jpeg", "image/png", "image/webp", "image/heic",
    "application/pdf",
}

# Word documents (.docx). Recognized separately from ALLOWED_TYPES because
# they need DIFFERENT handling — not uploaded-for-viewing-then-OCR'd like an
# image/PDF, but read for their actual text (see extract_docx_text below).
# Before this existed, a .docx attachment matched NO entry in ALLOWED_TYPES,
# so extract_attachments() silently dropped it entirely. With attachments then
# empty, the "no real attachment, render the HTML body to PDF" fallback further
# down fired on the WRAPPER email ("Attached please find the receipt...") and
# turned THAT into a synthetic PDF — the actual receipt content in the Word
# doc was never read at all. Old binary .doc (not .docx) isn't handled here;
# python-docx only reads the OOXML format.
DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

# python-docx reads a .docx's real text directly — no OCR needed at all, since
# unlike a scanned receipt this content is already machine-readable. Guarded
# import like weasyprint/PyMuPDF: a missing wheel degrades to today's
# behavior (the doc still uploads for the user to open by hand) rather than
# crashing the poller.
DOCX_AVAILABLE = False
DOCX_IMPORT_ERROR = None
try:
    import docx
    DOCX_AVAILABLE = True
except Exception as _ex:
    DOCX_IMPORT_ERROR = str(_ex)

def _is_docx_attachment(content_type: str, filename: str) -> bool:
    """True if this looks like a Word doc even when the sender mislabeled the
    MIME type (some mail clients send a generic application/octet-stream)."""
    if content_type in DOCX_CONTENT_TYPES:
        return True
    generic = {"application/octet-stream", "application/zip", ""}
    return content_type in generic and filename.lower().endswith(".docx")

# ── Logging ───────────────────────────────────────────────────────────────────
def log(msg: str):
    print(f"[{datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')}] {msg}", flush=True)

# ── IMAP helpers ──────────────────────────────────────────────────────────────
def fetch_unseen(imap: imaplib.IMAP4_SSL):
    """Return list of (uid_bytes, email.message.Message) for UNSEEN messages."""
    imap.select("INBOX")
    _, uid_data = imap.uid("search", None, "UNSEEN")
    uids = uid_data[0].split()
    messages = []
    for uid in uids:
        _, msg_data = imap.uid("fetch", uid, "(BODY.PEEK[])")
        raw = msg_data[0][1]
        msg = email.message_from_bytes(raw, policy=email.policy.default)
        messages.append((uid, msg))
    return messages

def mark_seen(imap: imaplib.IMAP4_SSL, uid: bytes):
    imap.uid("store", uid, "+FLAGS", "\\Seen")

def extract_attachments(msg):
    """Return list of {filename, content_type, data (base64 str)} for relevant parts.

    Accepts images regardless of disposition, and PDFs whether marked
    as inline or attachment — some vendors (DocuSign, Square, European
    e-invoice senders) ship the receipt PDF with `Content-Disposition: inline`
    and the previous attachment-only filter dropped them on the floor.
    """
    attachments = []
    for part in msg.walk():
        ct = part.get_content_type()
        filename = part.get_filename() or ""
        is_docx = _is_docx_attachment(ct, filename)
        if ct not in ALLOWED_TYPES and not is_docx:
            continue
        payload = part.get_payload(decode=True)
        if not payload or len(payload) > CONFIG["MAX_ATTACH_BYTES"]:
            continue
        # Normalize a mislabeled docx (e.g. application/octet-stream) to the
        # canonical MIME type so every downstream check (DOCX_CONTENT_TYPES
        # lookups here, and the edge function's own content-type handling)
        # agrees on what this file is.
        if is_docx:
            ct = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = filename or f"attachment.{ct.split('/')[-1]}"
        attachments.append({
            "filename": str(filename),
            "content_type": ct,
            "data": base64.b64encode(payload).decode("ascii"),
        })
    return attachments

def extract_docx_text(docx_bytes: bytes) -> str | None:
    """Extract plain text from a .docx (Word) attachment.

    Unlike a scanned receipt, this content is already machine-readable — no
    OCR needed, no rasterizing. python-docx reads the paragraphs and, since a
    contractor's invoice/receipt often puts the actual line items and total in
    a table rather than plain paragraphs, the table cells too. Returns None on
    any failure so the caller falls back to whatever the email body itself
    carried, same failure mode as the PDF/weasyprint helpers above.
    """
    if not DOCX_AVAILABLE:
        log(f"  extract_docx_text: python-docx not available (import error: {DOCX_IMPORT_ERROR})")
        return None
    try:
        document = docx.Document(io.BytesIO(docx_bytes))
        lines = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n".join(lines).strip()
        if text:
            log(f"  extract_docx_text: extracted {len(text)} chars")
        return text or None
    except Exception as ex:
        log(f"  extract_docx_text: FAILED ({type(ex).__name__}: {ex})")
        return None

def extract_body(msg):
    """Return (body_text, body_html) for the first text/plain and text/html parts.

    Many vendors (Uber, Lyft, airline/hotel "your receipt" emails, SaaS
    invoices, etc.) embed the receipt directly in the message body — no
    PDF or image attached. We forward both flavors when present so the
    edge function can OCR the inline content the same way it handles
    attached images.
    """
    body_text = None
    body_html = None
    for part in msg.walk():
        ct = part.get_content_type()
        cd = str(part.get("Content-Disposition", ""))
        # Skip anything explicitly attached — that's handled by extract_attachments.
        if "attachment" in cd:
            continue
        if ct == "text/plain" and body_text is None:
            payload = part.get_payload(decode=True)
            if payload:
                charset = part.get_content_charset() or "utf-8"
                try:
                    body_text = payload.decode(charset, errors="replace")
                except Exception:
                    body_text = payload.decode("utf-8", errors="replace")
        elif ct == "text/html" and body_html is None:
            payload = part.get_payload(decode=True)
            if payload:
                charset = part.get_content_charset() or "utf-8"
                try:
                    body_html = payload.decode(charset, errors="replace")
                except Exception:
                    body_html = payload.decode("utf-8", errors="replace")
    return body_text, body_html

def render_html_to_pdf(html: str) -> bytes | None:
    """Render HTML email body to PDF using weasyprint.

    Returns PDF bytes if successful, None otherwise.
    Used as a fallback when email has no real attachments but contains
    receipt/invoice content inline (e.g., Uber, Lyft, airline receipts).
    """
    if not WEASYPRINT_AVAILABLE:
        log(f"  render_html_to_pdf: weasyprint not available "
            f"(import error: {WEASYPRINT_IMPORT_ERROR})")
        return None
    try:
        # Wrap in basic styling. Block external network fetches so
        # missing remote images don't hang the render — weasyprint can
        # block on slow CDNs/tracking pixels for tens of seconds.
        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; margin: 20px; color: #333; }}
        img {{ max-width: 100%; height: auto; }}
        table {{ border-collapse: collapse; width: 100%; }}
        td, th {{ padding: 8px; border: 1px solid #ddd; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""

        # Custom URL fetcher that blocks all external requests — many
        # marketing emails contain tracking pixels and remote images
        # that would otherwise stall the render or fail noisily.
        def _no_network_fetcher(url):
            return {"string": b"", "mime_type": "image/png"}

        pdf_bytes = HTML(string=styled_html, url_fetcher=_no_network_fetcher).write_pdf()
        if pdf_bytes:
            log(f"  render_html_to_pdf: success ({len(pdf_bytes)} bytes)")
            return pdf_bytes
        log("  render_html_to_pdf: write_pdf returned empty")
        return None
    except Exception as ex:
        log(f"  render_html_to_pdf: FAILED ({type(ex).__name__}: {ex})")
        return None

def docx_text_to_preview_pdf(text: str) -> bytes | None:
    """Render extracted .docx text into a real PDF the browser can show
    inline, the same way it already shows an uploaded PDF.

    No browser has a native Word-document viewer, unlike PDF — the only ways
    to preview a raw .docx inline are Microsoft's or Google's online-viewer
    embeds, and both require sending the (signed) file URL to that third
    party to fetch and render, which isn't something to do with a household's
    receipt data without asking first. This stays in-house instead: reuse the
    weasyprint pipeline already here for HTML-body rendering, just fed the
    extracted text instead of an email body. The original .docx is still kept
    as its own attachment for anyone who wants the real editable file — this
    is purely an additional, clickable preview.
    """
    escaped_lines = [html.escape(line) if line.strip() else "" for line in text.split("\n")]
    body_html = "<br>\n".join(escaped_lines)
    fragment = (
        '<div style="font-size: 14px; line-height: 1.6; white-space: pre-wrap; '
        'font-family: Menlo, Consolas, monospace;">'
        f"{body_html}</div>"
    )
    return render_html_to_pdf(fragment)

# Rendering DPI for the OCR companion image. 150 keeps receipt text legible
# without producing a multi-megabyte PNG; the vision call downsamples anyway,
# so going higher buys nothing but upload time.
PDF_RASTER_DPI = 150

def rasterize_pdf_first_page(pdf_bytes: bytes) -> bytes | None:
    """Render page 1 of a PDF to PNG bytes so it can be OCR'd.

    Only page 1: receipts are effectively always single-page, and the edge
    function OCRs exactly one attachment, so extra pages would be uploaded
    and never read.

    Returns None on any failure — callers fall back to today's behaviour of
    sending the PDF alone, which is no worse than before.
    """
    if not PYMUPDF_AVAILABLE:
        log(f"  rasterize_pdf: PyMuPDF not available "
            f"(import error: {PYMUPDF_IMPORT_ERROR})")
        return None
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        if doc.page_count == 0:
            log("  rasterize_pdf: PDF has no pages")
            doc.close()
            return None
        pix = doc.load_page(0).get_pixmap(dpi=PDF_RASTER_DPI)
        png_bytes = pix.tobytes("png")
        doc.close()

        if len(png_bytes) > CONFIG["MAX_ATTACH_BYTES"]:
            log(f"  rasterize_pdf: PNG too large ({len(png_bytes)} bytes), skipping")
            return None
        log(f"  rasterize_pdf: page 1 -> PNG ({len(png_bytes)} bytes @ {PDF_RASTER_DPI}dpi)")
        return png_bytes
    except Exception as ex:
        log(f"  rasterize_pdf: FAILED ({type(ex).__name__}: {ex})")
        return None

def add_ocr_companion_images(attachments: list) -> list:
    """Insert a PNG render ahead of each PDF so OCR has something it can read.

    The edge function scans for the FIRST OCR-compatible attachment, so
    position matters: putting the PNG before its PDF means even the currently
    deployed function (which still tries, and fails, to OCR PDFs directly)
    picks up the image instead. The original PDF is kept and still uploaded —
    it remains the better artifact for a human to open.
    """
    if not any(a["content_type"] == "application/pdf" for a in attachments):
        return attachments

    result = []
    rendered = 0
    for att in attachments:
        if att["content_type"] == "application/pdf" and rendered == 0:
            try:
                pdf_bytes = base64.b64decode(att["data"])
            except Exception as ex:
                log(f"  rasterize_pdf: could not decode {att['filename']} ({ex})")
                result.append(att)
                continue
            png_bytes = rasterize_pdf_first_page(pdf_bytes)
            if png_bytes:
                stem = att["filename"].rsplit(".", 1)[0] or "attachment"
                result.append({
                    "filename": f"{stem}.png",
                    "content_type": "image/png",
                    "data": base64.b64encode(png_bytes).decode("ascii"),
                })
                rendered += 1
        result.append(att)
    return result

# ── Edge function call ────────────────────────────────────────────────────────
def post_to_function(payload: dict) -> bool:
    body = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        CONFIG["FUNCTION_URL"],
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {CONFIG['INBOUND_SECRET']}",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            result = json.loads(resp.read())
            log(f"  → function response: {result}")
            return True
    except urllib.error.HTTPError as e:
        log(f"  → HTTP {e.code}: {e.read().decode()}")
        return False
    except Exception as ex:
        log(f"  → error calling function: {ex}")
        return False

# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    if CONFIG["IMAP_PASSWORD"] == "REPLACE_WITH_MAILBOX_PASSWORD":
        log("ERROR: IMAP password not configured. Set LEDGERX_IMAP_PASSWORD env var.")
        sys.exit(1)
    if CONFIG["INBOUND_SECRET"] == "REPLACE_WITH_SHARED_SECRET":
        log("ERROR: Inbound secret not configured. Set LEDGERX_INBOUND_SECRET env var.")
        sys.exit(1)

    if WEASYPRINT_AVAILABLE:
        log(f"weasyprint loaded: version={WEASYPRINT_VERSION}")
    else:
        log(f"weasyprint NOT available: {WEASYPRINT_IMPORT_ERROR}")

    log(f"Connecting to {CONFIG['IMAP_HOST']}:{CONFIG['IMAP_PORT']} as {CONFIG['IMAP_USER']}")
    try:
        imap = imaplib.IMAP4_SSL(CONFIG["IMAP_HOST"], CONFIG["IMAP_PORT"])
        imap.login(CONFIG["IMAP_USER"], CONFIG["IMAP_PASSWORD"])
    except Exception as ex:
        log(f"IMAP connection failed: {ex}")
        sys.exit(1)

    messages = fetch_unseen(imap)
    log(f"Found {len(messages)} unseen message(s)")

    for uid, msg in messages:
        # `policy.default` can return Header objects for encoded values
        # (e.g. RFC2047-encoded non-ASCII subjects). Coerce to plain str
        # so json.dumps in post_to_function() doesn't choke and so
        # downstream string ops are safe.
        from_addr = str(msg.get("From", "") or "")
        subject   = str(msg.get("Subject", "") or "")
        # Strip angle brackets so the edge function's dedup lookup
        # matches consistently regardless of which mail client added them.
        msg_id    = str(msg.get("Message-ID", "") or "").strip().strip("<>").strip()

        # Extract just the email address from "Name <addr>" format
        import re
        m = re.search(r"<([^>]+)>", from_addr)
        from_email = m.group(1).strip() if m else from_addr.strip()

        log(f"Processing: from={from_email!r} subject={subject[:80]!r} msg_id={msg_id!r}")

        attachments = extract_attachments(msg)
        log(f"  Attachments: {[a['filename'] for a in attachments]}")

        body_text, body_html = extract_body(msg)
        log(f"  Body: text={'yes' if body_text else 'no'} html={'yes' if body_html else 'no'}")

        # Word-doc attachments: read their real text directly (no OCR needed —
        # it's already machine-readable) and feed it through the same inline-
        # body extraction path a plain-text email uses. This ALSO prevents the
        # "no real attachment" HTML-to-PDF fallback just below from firing on
        # a docx-only forward and turning the wrapper email text into a
        # synthetic PDF while the actual receipt content sits unread in the
        # attachment — attachments is non-empty now that the docx is kept, so
        # that branch correctly no longer applies here at all.
        # ALSO render that same text into a real PDF via the existing
        # weasyprint pipeline and insert it just ahead of the original .docx,
        # so clicking the attachment in the app opens an in-browser preview
        # exactly like an uploaded PDF does. No browser has a native .docx
        # renderer — the only alternatives (Microsoft/Google's online-viewer
        # embeds) require sending the file's URL to that third party to fetch
        # and render, which isn't something to do with a household's receipt
        # data without asking first. The real .docx is kept too, right after
        # its preview, for anyone who wants the original editable file.
        docx_texts = []
        attachments_with_previews = []
        for att in attachments:
            if att["content_type"] in DOCX_CONTENT_TYPES:
                text = extract_docx_text(base64.b64decode(att["data"]))
                if text:
                    docx_texts.append(text)
                    preview_bytes = docx_text_to_preview_pdf(text)
                    if preview_bytes and len(preview_bytes) < CONFIG["MAX_ATTACH_BYTES"]:
                        stem = att["filename"].rsplit(".", 1)[0] or "receipt"
                        attachments_with_previews.append({
                            "filename": f"{stem}.pdf",
                            "content_type": "application/pdf",
                            "data": base64.b64encode(preview_bytes).decode("ascii"),
                        })
                        log(f"  Rendered Word-doc preview PDF ({len(preview_bytes)} bytes)")
            attachments_with_previews.append(att)
        attachments = attachments_with_previews

        if docx_texts:
            combined = "\n\n".join(docx_texts)
            body_text = f"{combined}\n\n{body_text}" if body_text else combined
            log(f"  Word attachment text merged into body ({len(combined)} chars from {len(docx_texts)} file(s))")

        # If no real attachments but we have HTML body, render it to PDF
        # instead of sending raw body text for OCR (user will review visually).
        if not attachments and body_html:
            pdf_bytes = render_html_to_pdf(body_html)
            if pdf_bytes and len(pdf_bytes) < CONFIG["MAX_ATTACH_BYTES"]:
                attachments.append({
                    "filename": "email-body.pdf",
                    "content_type": "application/pdf",
                    "data": base64.b64encode(pdf_bytes).decode("ascii"),
                })
                log(f"  Rendered HTML body to PDF ({len(pdf_bytes)} bytes)")
                # Deliberately KEEP body_text/body_html. This used to null them
                # out on the reasoning that the PDF superseded them — but the
                # edge function can't OCR a PDF, so that left it with nothing at
                # all to read and every such receipt landed with prefilled={}.
                # The inline-body extractor only runs when attachment OCR came
                # back empty, so keeping these costs nothing in the happy path
                # and is the difference between a prefilled form and a blank one
                # if rasterizing is unavailable.
            else:
                log(f"  HTML to PDF rendering failed or output too large, sending body as text")

        # Give every PDF an OCR-readable PNG companion, inserted ahead of it.
        attachments = add_ocr_companion_images(attachments)

        payload = {
            "from_email": from_email,
            "subject": subject,
            "message_id": msg_id or None,
            "attachments": attachments,
            "body_text": body_text,
            "body_html": body_html,
        }

        success = post_to_function(payload)
        if success:
            mark_seen(imap, uid)
            log(f"  Marked as seen.")
        else:
            log(f"  Left as unseen (will retry next poll).")

    imap.logout()
    log("Done.")

if __name__ == "__main__":
    main()
